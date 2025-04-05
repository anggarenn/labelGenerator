import streamlit as st
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import os

def create_label_docx(daftar_nama, template_awal, template_akhir, nama_output="label_undangan.docx"):
    # Jika template kosong, gunakan default
    if not template_awal:
        template_awal = "Kepada Yth,"
    if not template_akhir:
        template_akhir = "Di Tempat."

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)

    label_per_halaman = 12
    jumlah_halaman = (len(daftar_nama) + label_per_halaman - 1) // label_per_halaman

    index = 0
    for _ in range(jumlah_halaman):
        table = doc.add_table(rows=4, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        for row in table.rows:
            row.height = Cm(3.2)
            row.height_rule = True

        for r in range(4):
            for c in range(3):
                if index >= len(daftar_nama):
                    continue

                cell = table.cell(r, c)
                cell.width = Cm(6.4)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.left_indent = Cm(1.2)
                p.paragraph_format.line_spacing = Pt(12)

                if r == 0:
                    space_top = Pt(40)
                elif r == 1:
                    space_top = Pt(65)
                elif r == 2:
                    space_top = Pt(70)
                elif r == 3:
                    space_top = Pt(60)

                p.paragraph_format.space_before = space_top

                # Gunakan template yang diinput atau default
                run = p.add_run(f"{template_awal}\n{daftar_nama[index]}\n{template_akhir}")
                run.font.name = "Calibri"
                run.font.size = Pt(11)

                index += 1

        if index < len(daftar_nama):
            doc.add_page_break()

    # Simpan hasil file
    doc.save(nama_output)
    return nama_output

# Streamlit Web App
st.title("Generator Label Undangan")
st.write("Pilih opsi di bawah untuk membuat label undangan:")

# Menambahkan watermark di halaman Streamlit
st.markdown(
    """
    <style>
        .watermark {
            position: absolute;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%);
            font-size: 70px;
            color: rgba(0, 0, 0, 0.1);
            z-index: -1;
            font-weight: bold;
            pointer-events: none;
            white-space: nowrap;
        }
    </style>
    <div class="watermark">
        WATERMARK
    </div>
    """, unsafe_allow_html=True)

# Pilih opsi input manual atau upload file
input_option = st.radio("Pilih cara input daftar nama", ("Input Manual", "Upload File .txt"))

# Input template yang bisa diedit
st.write("Anda dapat mengubah template untuk bagian salam dan alamat:")
template_awal = st.text_input("Template Salam (misalnya: 'Kepada Yth,')", "Kepada Yth,")
template_akhir = st.text_input("Template Alamat (misalnya: 'Di Tempat.')", "Di Tempat.")

if input_option == "Input Manual":
    # Input manual daftar nama
    daftar_nama_input = st.text_area("Masukkan daftar nama (pisahkan setiap nama dengan baris baru)")
    
    # Validasi input
    if st.button("Generate Label"):
        if not daftar_nama_input.strip():  # Jika input kosong
            st.error("❌ Data daftar nama tamu tidak boleh kosong.")
        else:
            daftar_nama = daftar_nama_input.splitlines()
            output_file = "label_undangan.docx"
            result = create_label_docx(daftar_nama, template_awal, template_akhir, output_file)
            st.success(f"✅ Label berhasil dibuat! File disimpan di: {output_file}")

            # Offer untuk download file
            with open(output_file, "rb") as f:
                st.download_button("Download Label", f, file_name=output_file, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

elif input_option == "Upload File .txt":
    # Upload file .txt
    uploaded_file = st.file_uploader("Pilih file daftar nama", type="txt")

    if uploaded_file is not None:
        # Simpan file sementara
        temp_filename = "uploaded_daftar_nama.txt"
        with open(temp_filename, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        st.write(f"File {uploaded_file.name} berhasil diupload!")

        # Button untuk generate label
        if st.button("Generate Label"):
            with open(temp_filename, encoding='utf-8') as f:
                daftar_nama = [line.strip() for line in f if line.strip()]

            # Validasi input
            if not daftar_nama:
                st.error("❌ Data daftar nama tamu tidak boleh kosong.")
            else:
                output_file = "label_undangan.docx"
                result = create_label_docx(daftar_nama, template_awal, template_akhir, output_file)
                st.success(f"✅ Label berhasil dibuat! File disimpan di: {output_file}")

                # Offer untuk download file
                with open(output_file, "rb") as f:
                    st.download_button("Download Label", f, file_name=output_file, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
        # Delete temporary file
        os.remove(temp_filename)
